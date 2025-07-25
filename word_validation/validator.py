from typing import List, Dict, Any, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from word_formatting.manager import WordFormattingManager, WordRunFormatting
from word_processors.text_processor import WordTextProcessor
import logging

logger = logging.getLogger('word_translator')

class WordFormatValidator:
    """Validator for ensuring Word document formatting preservation"""
    
    def __init__(self, formatting_manager: WordFormattingManager):
        self.formatting_manager = formatting_manager
        self.text_processor = WordTextProcessor()
        
    def validate_document_structure(self, original_path: str, translated_path: str) -> Dict[str, Any]:
        """
        Validate that the translated document maintains the same structure as the original.
        Returns a comprehensive validation report.
        """
        try:
            original_doc = Document(original_path)
            translated_doc = Document(translated_path)
            
            validation_report = {
                'structure_match': True,
                'warnings': [],
                'errors': [],
                'statistics': {
                    'original_paragraphs': len(original_doc.paragraphs),
                    'translated_paragraphs': len(translated_doc.paragraphs),
                    'original_tables': len(original_doc.tables),
                    'translated_tables': len(translated_doc.tables),
                    'formatting_preservation_score': 0.0
                }
            }
            
            # Validate paragraph count
            if len(original_doc.paragraphs) != len(translated_doc.paragraphs):
                validation_report['structure_match'] = False
                validation_report['errors'].append(
                    f"Paragraph count mismatch: {len(original_doc.paragraphs)} vs {len(translated_doc.paragraphs)}"
                )
            
            # Validate table count
            if len(original_doc.tables) != len(translated_doc.tables):
                validation_report['structure_match'] = False
                validation_report['errors'].append(
                    f"Table count mismatch: {len(original_doc.tables)} vs {len(translated_doc.tables)}"
                )
            
            # Validate paragraph-level formatting
            formatting_score = self._validate_paragraph_formatting(
                original_doc.paragraphs, 
                translated_doc.paragraphs,
                validation_report
            )
            validation_report['statistics']['formatting_preservation_score'] = formatting_score
            
            # Validate table formatting
            if original_doc.tables and translated_doc.tables:
                self._validate_table_formatting(
                    original_doc.tables,
                    translated_doc.tables,
                    validation_report
                )
            
            return validation_report
            
        except Exception as e:
            logger.error(f"Error during validation: {str(e)}")
            return {
                'structure_match': False,
                'warnings': [],
                'errors': [f"Validation failed: {str(e)}"],
                'statistics': {}
            }
    
    def _validate_paragraph_formatting(self, original_paras: List[Paragraph], 
                                     translated_paras: List[Paragraph],
                                     validation_report: Dict[str, Any]) -> float:
        """Validate formatting preservation at paragraph level"""
        formatting_matches = 0
        total_comparisons = 0
        
        for orig_para, trans_para in zip(original_paras, translated_paras):
            if orig_para.text.strip() and trans_para.text.strip():
                total_comparisons += 1
                
                # Compare run counts
                if len(orig_para.runs) != len(trans_para.runs):
                    validation_report['warnings'].append(
                        f"Run count mismatch in paragraph: {len(orig_para.runs)} vs {len(trans_para.runs)}"
                    )
                else:
                    formatting_matches += 1
                
                # Compare run-level formatting
                for orig_run, trans_run in zip(orig_para.runs, trans_para.runs):
                    if not self._compare_run_formatting(orig_run, trans_run):
                        validation_report['warnings'].append(
                            "Run formatting mismatch detected"
                        )
        
        return formatting_matches / total_comparisons if total_comparisons > 0 else 1.0
    
    def _validate_table_formatting(self, original_tables: List, translated_tables: List,
                                 validation_report: Dict[str, Any]) -> None:
        """Validate table structure and formatting preservation"""
        for table_idx, (orig_table, trans_table) in enumerate(zip(original_tables, translated_tables)):
            # Check row count
            if len(orig_table.rows) != len(trans_table.rows):
                validation_report['errors'].append(
                    f"Table {table_idx + 1} row count mismatch: "
                    f"{len(orig_table.rows)} vs {len(trans_table.rows)}"
                )
            
            # Check column count for each row
            for row_idx, (orig_row, trans_row) in enumerate(zip(orig_table.rows, trans_table.rows)):
                if len(orig_row.cells) != len(trans_row.cells):
                    validation_report['errors'].append(
                        f"Table {table_idx + 1} row {row_idx + 1} column count mismatch: "
                        f"{len(orig_row.cells)} vs {len(trans_row.cells)}"
                    )
    
    def _compare_run_formatting(self, original_run: Run, translated_run: Run) -> bool:
        """Compare formatting between two runs"""
        try:
            orig_format = self.formatting_manager.collect_run_formatting(original_run)
            trans_format = self.formatting_manager.collect_run_formatting(translated_run)
            
            # Compare key formatting attributes
            return (
                orig_format.bold == trans_format.bold and
                orig_format.italic == trans_format.italic and
                orig_format.underline == trans_format.underline and
                orig_format.font_name == trans_format.font_name
            )
        except Exception as e:
            logger.warning(f"Error comparing run formatting: {str(e)}")
            return False
    
    def validate_translation_quality(self, original_path: str, translated_path: str) -> Dict[str, Any]:
        """
        Validate translation quality by checking for common issues.
        """
        try:
            original_doc = Document(original_path)
            translated_doc = Document(translated_path)
            
            quality_report = {
                'quality_score': 0.0,
                'issues': [],
                'statistics': {
                    'unchanged_paragraphs': 0,
                    'total_paragraphs': 0,
                    'empty_translations': 0
                }
            }
            
            total_paragraphs = 0
            unchanged_count = 0
            empty_translations = 0
            
            for orig_para, trans_para in zip(original_doc.paragraphs, translated_doc.paragraphs):
                orig_text = orig_para.text.strip()
                trans_text = trans_para.text.strip()
                
                if orig_text:
                    total_paragraphs += 1
                    
                    # Check for unchanged text (potential translation failure)
                    if orig_text == trans_text:
                        unchanged_count += 1
                        quality_report['issues'].append(f"Unchanged text: {orig_text[:50]}...")
                    
                    # Check for empty translations
                    if orig_text and not trans_text:
                        empty_translations += 1
                        quality_report['issues'].append(f"Empty translation for: {orig_text[:50]}...")
            
            quality_report['statistics']['total_paragraphs'] = total_paragraphs
            quality_report['statistics']['unchanged_paragraphs'] = unchanged_count
            quality_report['statistics']['empty_translations'] = empty_translations
            
            # Calculate quality score
            if total_paragraphs > 0:
                success_rate = (total_paragraphs - unchanged_count - empty_translations) / total_paragraphs
                quality_report['quality_score'] = success_rate
            else:
                quality_report['quality_score'] = 1.0
            
            return quality_report
            
        except Exception as e:
            logger.error(f"Error during quality validation: {str(e)}")
            return {
                'quality_score': 0.0,
                'issues': [f"Quality validation failed: {str(e)}"],
                'statistics': {}
            }

class WordVisualComparator:
    """Visual comparison utilities for Word documents"""
    
    def __init__(self, formatting_manager: WordFormattingManager):
        self.formatting_manager = formatting_manager
    
    def generate_formatting_report(self, document_path: str) -> Dict[str, Any]:
        """Generate a detailed formatting report for a Word document"""
        try:
            doc = Document(document_path)
            
            report = {
                'document_path': document_path,
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'formatting_analysis': [],
                'style_usage': {},
                'font_usage': {},
                'complex_formatting_count': 0
            }
            
            # Analyze paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_analysis = self._analyze_paragraph_formatting(paragraph, para_idx)
                    report['formatting_analysis'].append(para_analysis)
                    
                    # Count complex formatting
                    if para_analysis['run_count'] > 1:
                        report['complex_formatting_count'] += 1
                    
                    # Track style usage
                    style_name = paragraph.style.name if paragraph.style else "None"
                    report['style_usage'][style_name] = report['style_usage'].get(style_name, 0) + 1
                    
                    # Track font usage
                    for run in paragraph.runs:
                        font_name = run.font.name or "Default"
                        report['font_usage'][font_name] = report['font_usage'].get(font_name, 0) + 1
            
            return report
            
        except Exception as e:
            logger.error(f"Error generating formatting report: {str(e)}")
            return {'error': str(e)}
    
    def _analyze_paragraph_formatting(self, paragraph: Paragraph, para_idx: int) -> Dict[str, Any]:
        """Analyze formatting of a single paragraph"""
        analysis = {
            'paragraph_index': para_idx,
            'text_preview': paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
            'run_count': len(paragraph.runs),
            'style': paragraph.style.name if paragraph.style else None,
            'runs': []
        }
        
        for run_idx, run in enumerate(paragraph.runs):
            run_analysis = {
                'run_index': run_idx,
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline
            }
            analysis['runs'].append(run_analysis)
        
        return analysis
    
    def compare_documents(self, original_path: str, translated_path: str) -> Dict[str, Any]:
        """Compare two documents and generate a detailed comparison report"""
        try:
            original_report = self.generate_formatting_report(original_path)
            translated_report = self.generate_formatting_report(translated_path)
            
            comparison = {
                'original_document': original_report,
                'translated_document': translated_report,
                'comparison_summary': {
                    'structure_preserved': True,
                    'formatting_preservation_score': 0.0,
                    'differences': []
                }
            }
            
            # Compare structure
            if (original_report.get('total_paragraphs', 0) != 
                translated_report.get('total_paragraphs', 0)):
                comparison['comparison_summary']['structure_preserved'] = False
                comparison['comparison_summary']['differences'].append(
                    f"Paragraph count difference: {original_report.get('total_paragraphs', 0)} vs "
                    f"{translated_report.get('total_paragraphs', 0)}"
                )
            
            # Compare complex formatting preservation
            orig_complex = original_report.get('complex_formatting_count', 0)
            trans_complex = translated_report.get('complex_formatting_count', 0)
            
            if orig_complex > 0:
                formatting_score = trans_complex / orig_complex
                comparison['comparison_summary']['formatting_preservation_score'] = formatting_score
                
                if formatting_score < 1.0:
                    comparison['comparison_summary']['differences'].append(
                        f"Complex formatting loss: {orig_complex} vs {trans_complex} paragraphs"
                    )
            else:
                comparison['comparison_summary']['formatting_preservation_score'] = 1.0
            
            return comparison
            
        except Exception as e:
            logger.error(f"Error comparing documents: {str(e)}")
            return {'error': str(e)}