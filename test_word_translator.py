#!/usr/bin/env python3
"""
Test script for Word Document Translator
Demonstrates the formatting preservation capabilities and provides validation.
"""

import os
import sys
from typing import Callable
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from translate_word import WordTranslator, WORD_INSTRUCTIONS
from word_validation.validator import WordFormatValidator, WordVisualComparator
from word_formatting.manager import WordFormattingManager
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger('word_translator_test')

def create_test_document(path: str) -> None:
    """Create a complex test Word document with various formatting"""
    doc = Document()
    
    # Title with custom formatting
    title = doc.add_heading('Test Document for Translation', 0)
    title_run = title.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    
    # Paragraph with mixed formatting
    para1 = doc.add_paragraph()
    run1 = para1.add_run('This is a ')
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    
    run2 = para1.add_run('bold text')
    run2.bold = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    
    run3 = para1.add_run(' and this is ')
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(12)
    
    run4 = para1.add_run('italic text')
    run4.italic = True
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(12)
    
    run5 = para1.add_run(' with ')
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(12)
    
    run6 = para1.add_run('underlined content')
    run6.underline = True
    run6.font.name = 'Times New Roman'
    run6.font.size = Pt(12)
    run6.font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    run7 = para1.add_run('.')
    run7.font.name = 'Times New Roman'
    run7.font.size = Pt(12)
    
    # Paragraph with different alignment
    para2 = doc.add_paragraph('This paragraph is center-aligned with custom spacing.')
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_before = Pt(12)
    para2.paragraph_format.space_after = Pt(12)
    
    # Add a table with formatting
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Product'
    header_cells[1].text = 'Quantity'
    header_cells[2].text = 'Price'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    data_rows = [
        ['Laptop Computer', '5 units', '$2,500.00'],
        ['Office Chair', '10 units', '$150.00']
    ]
    
    for row_idx, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            if col_idx == 2:  # Price column - make it bold and right-aligned
                for paragraph in row_cells[col_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    
    # Paragraph with special characters and formatting
    para3 = doc.add_paragraph()
    run1 = para3.add_run('Special formatting: ')
    run1.font.size = Pt(14)
    
    run2 = para3.add_run('ACRONYM')
    run2.font.size = Pt(14)
    run2.all_caps = True
    
    run3 = para3.add_run(' and ')
    run3.font.size = Pt(14)
    
    run4 = para3.add_run('subscript')
    run4.font.size = Pt(14)
    run4.subscript = True
    
    run5 = para3.add_run(' and ')
    run5.font.size = Pt(14)
    
    run6 = para3.add_run('superscript')
    run6.font.size = Pt(14)
    run6.superscript = True
    
    run7 = para3.add_run(' text.')
    run7.font.size = Pt(14)
    
    # Save the document
    doc.save(path)
    logger.info(f"Created test document: {path}")

def mock_translate_function(text: str) -> str:
    """Mock translation function for testing - adds [TRANSLATED] prefix"""
    if not text or not text.strip():
        return text
    
    # Simple mock translation that preserves structure
    lines = text.split('\\n')
    translated_lines = []
    
    for line in lines:
        if line.strip():
            # Add [TRANSLATED] prefix to show translation occurred
            translated_line = f"[TRANSLATED] {line}"
            translated_lines.append(translated_line)
        else:
            translated_lines.append(line)
    
    return '\\n'.join(translated_lines)

def test_word_translator():
    """Test the Word translator with a complex document"""
    # Paths
    test_doc_path = "test_document.docx"
    translated_doc_path = "test_document_translated.docx"
    
    try:
        # Create test document
        logger.info("Creating test document...")
        create_test_document(test_doc_path)
        
        # Initialize components
        formatting_manager = WordFormattingManager()
        validator = WordFormatValidator(formatting_manager)
        comparator = WordVisualComparator(formatting_manager)
        
        # Generate original document report
        logger.info("Analyzing original document...")
        original_report = comparator.generate_formatting_report(test_doc_path)
        logger.info(f"Original document has {original_report['total_paragraphs']} paragraphs")
        logger.info(f"Complex formatting in {original_report['complex_formatting_count']} paragraphs")
        
        # Create translator (using mock function to avoid API calls)
        logger.info("Initializing translator...")
        translator = WordTranslator(
            api_key="test-key",  # Not used with mock function
            model="gpt-4o-mini",
            translate_headers=True,
            translate_footers=True,
            translate_comments=False
        )
        
        # Translate document using mock function
        logger.info("Translating document...")
        translator.translate_document(
            test_doc_path, 
            translated_doc_path, 
            translation_func=mock_translate_function
        )
        
        # Validate translation
        logger.info("Validating translation...")
        structure_validation = validator.validate_document_structure(
            test_doc_path, 
            translated_doc_path
        )
        
        quality_validation = validator.validate_translation_quality(
            test_doc_path,
            translated_doc_path
        )
        
        # Generate comparison report
        comparison_report = comparator.compare_documents(
            test_doc_path,
            translated_doc_path
        )
        
        # Print results
        print("\\n" + "="*80)
        print("WORD TRANSLATOR TEST RESULTS")
        print("="*80)
        
        print(f"\\nStructure Validation:")
        print(f"  Structure Match: {structure_validation['structure_match']}")
        print(f"  Formatting Score: {structure_validation['statistics']['formatting_preservation_score']:.2%}")
        print(f"  Errors: {len(structure_validation['errors'])}")
        print(f"  Warnings: {len(structure_validation['warnings'])}")
        
        if structure_validation['errors']:
            print("  Errors:")
            for error in structure_validation['errors']:
                print(f"    - {error}")
        
        if structure_validation['warnings']:
            print("  Warnings:")
            for warning in structure_validation['warnings'][:5]:  # Show first 5
                print(f"    - {warning}")
        
        print(f"\\nTranslation Quality:")
        print(f"  Quality Score: {quality_validation['quality_score']:.2%}")
        print(f"  Total Paragraphs: {quality_validation['statistics']['total_paragraphs']}")
        print(f"  Unchanged: {quality_validation['statistics']['unchanged_paragraphs']}")
        print(f"  Empty Translations: {quality_validation['statistics']['empty_translations']}")
        
        print(f"\\nFormatting Preservation:")
        comparison = comparison_report['comparison_summary']
        print(f"  Structure Preserved: {comparison['structure_preserved']}")
        print(f"  Formatting Score: {comparison['formatting_preservation_score']:.2%}")
        
        if comparison['differences']:
            print("  Differences:")
            for diff in comparison['differences']:
                print(f"    - {diff}")
        
        # Print success message
        if (structure_validation['structure_match'] and 
            quality_validation['quality_score'] > 0.8 and
            comparison['formatting_preservation_score'] > 0.9):
            print(f"\\n✅ TEST PASSED - Word translator successfully preserved formatting!")
        else:
            print(f"\\n⚠️  TEST ISSUES - Some formatting or translation issues detected.")
        
        print(f"\\nFiles created:")
        print(f"  Original: {test_doc_path}")
        print(f"  Translated: {translated_doc_path}")
        print("="*80)
        
        return True
        
    except Exception as e:
        logger.error(f"Test failed: {str(e)}")
        print(f"\\n❌ TEST FAILED: {str(e)}")
        return False
    
    finally:
        # Clean up test files
        for path in [test_doc_path, translated_doc_path]:
            if os.path.exists(path):
                try:
                    os.remove(path)
                    logger.info(f"Cleaned up: {path}")
                except Exception as e:
                    logger.warning(f"Could not remove {path}: {str(e)}")

if __name__ == "__main__":
    print("Word Document Translator Test")
    print("="*40)
    
    # Test the Word translator
    success = test_word_translator()
    
    if success:
        print("\\nTest completed successfully!")
        sys.exit(0)
    else:
        print("\\nTest failed!")
        sys.exit(1)