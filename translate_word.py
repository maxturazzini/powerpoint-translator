"""
Word Document Translator V1
Implements comprehensive text formatting preservation during translation.

Features:
- Complete formatting preservation at run level
- Enhanced text processing for complex documents
- Table and header/footer support
- Comprehensive validation
"""

from openai import OpenAI
import logging
import sys
from typing import Optional, Callable
from docx import Document
from docx.shared import RGBColor
from word_formatting.manager import WordFormattingManager
from word_processors.text_processor import WordTextProcessor
from word_processors.enhanced_document_processor import EnhancedDocumentProcessor
from word_validation.validator import WordFormatValidator, WordVisualComparator

# Configure logging to both file and console
logger = logging.getLogger('word_translator')
logger.setLevel(logging.INFO)

# Clear any existing handlers
logger.handlers = []

# File handler
file_handler = logging.FileHandler('word_translator.log')
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(
    logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
)
logger.addHandler(file_handler)

# Console handler
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(
    logging.Formatter('%(message)s')  # Simplified console format
)
logger.addHandler(console_handler)

# Translation instructions for Word documents
WORD_INSTRUCTIONS = """
You are a professional translator specialising in business- and technical-oriented Word documents.
Your assignment is to translate all text written in Italian (SOURCE_LANG) into English (TARGET_LANG).

IMPORTANT: You must ALWAYS translate any Italian text to English. Never leave Italian text untranslated.

### Task
* Translate every segment of text that is in Italian to English taking into account the context of the text (i.e. technical terms, names, numbers, etc.).

### Translation guidelines

1. ALWAYS translate Italian text to clear, professional business English.
2. Maintain the original meaning and context.
3. Keep technical terms, names, and numbers unchanged when appropriate.
4. If text is already in English, return it unchanged.
5. Retain **all** original formatting — line breaks, spacing, punctuation, and special characters.
6. Provide **only** the translated text; add **no explanations or comments**.
7. Ensure natural, fluent English phrasing.
8. Use consistent terminology throughout.
9. Preserve line breaks and spacing exactly as in the original text.

### Example
**Input (Italian)**
Benvenuti al
documento finale.
**Output (English)**
Welcome to the
final document.

Remember: ALL Italian text MUST be translated to English.
"""

class WordTranslator:
    """Word document translator with comprehensive formatting preservation"""
    
    def __init__(
        self,
        api_key: str,
        model: str = "gpt-4o",  # upgraded for better translation quality
        translate_headers: bool = True,
        translate_footers: bool = True,
        translate_comments: bool = False
    ):
        """
        Initialize the translator with OpenAI credentials and options.
        
        Args:
            api_key: OpenAI API key
            model: OpenAI model to use
            translate_headers: Whether to translate headers
            translate_footers: Whether to translate footers
            translate_comments: Whether to translate comments
        """
        self.model = model
        self.client = OpenAI(api_key=api_key)
        self.translate_headers = translate_headers
        self.translate_footers = translate_footers
        self.translate_comments = translate_comments
        logger.info(f"Initialized Word translator with model: {model}")
        logger.info(f"Translate headers: {translate_headers}")
        logger.info(f"Translate footers: {translate_footers}")
        logger.info(f"Translate comments: {translate_comments}")
        
        # Initialize components
        self.formatting_manager = WordFormattingManager()
        self.enhanced_document_processor = EnhancedDocumentProcessor(self.formatting_manager)
        self.validator = WordFormatValidator(self.formatting_manager)
        
    def translate_text(self, text: str) -> str:
        """Translate text using OpenAI's API"""
        if not text or not text.strip():
            return text
            
        try:
            logger.info("-" * 80)
            logger.info(f"Original text: {text}")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": WORD_INSTRUCTIONS},
                    {"role": "user", "content": f"Translate this Italian text to English, preserving all line breaks and spacing:\n\n{text}"}
                ],
                temperature=0.1,  # Even lower temperature for more consistent translations
                max_tokens=1500,  # Increased token limit for longer translations
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
            
            translated = response.choices[0].message.content
            logger.info(f"Translated text: {translated}")
            logger.info("-" * 80)
            
            # Verify translation happened
            if translated.strip() == text.strip():
                logger.warning("WARNING: Text remained unchanged after translation attempt!")
            
            return translated
            
        except Exception as e:
            logger.error(f"Translation error: {str(e)}")
            return text
            
    def translate_document(
        self,
        input_path: str,
        output_path: str,
        translation_func: Optional[Callable[[str], str]] = None
    ) -> None:
        """
        Translate Word document while preserving formatting.
        
        Args:
            input_path: Path to the input Word file
            output_path: Path where the translated file will be saved
            translation_func: Optional custom translation function
        """
        # Use provided translation function or default to OpenAI
        translate = translation_func or self.translate_text
        
        try:
            # Load document
            logger.info(f"Loading document: {input_path}")
            document = Document(input_path)
            
            logger.info("\nStarting translation process...")
            translation_count = 0
            unchanged_count = 0
            
            # Process main document paragraphs
            logger.info("Processing main document paragraphs...")
            for para_idx, paragraph in enumerate(document.paragraphs, 1):
                if paragraph.text.strip():
                    logger.info(f"Processing paragraph {para_idx}")
                    original_text = paragraph.text.strip()
                    logger.info(f"Found text to translate: {original_text}")
                    
                    # Process paragraph with enhanced formatting preservation
                    self.enhanced_document_processor.process_paragraph(paragraph, translate)
                    
                    # Verify translation after processing
                    translated_text = paragraph.text.strip()
                    if translated_text:
                        if translated_text == original_text:
                            unchanged_count += 1
                            logger.warning(f"WARNING: Text remained unchanged: {original_text}")
                        else:
                            translation_count += 1
                            logger.info(f"Successfully translated: {original_text} -> {translated_text}")
            
            # Process tables
            logger.info("Processing tables...")
            for table_idx, table in enumerate(document.tables, 1):
                logger.info(f"Processing table {table_idx}")
                for row_idx, row in enumerate(table.rows, 1):
                    for cell_idx, cell in enumerate(row.cells, 1):
                        for para in cell.paragraphs:
                            if para.text.strip():
                                original_text = para.text.strip()
                                logger.info(f"Found table text to translate: {original_text}")
                                
                                # Process cell paragraph
                                self.enhanced_document_processor.process_paragraph(para, translate)
                                
                                translated_text = para.text.strip()
                                if translated_text != original_text:
                                    translation_count += 1
                                else:
                                    unchanged_count += 1
            
            # Process headers if configured
            if self.translate_headers:
                logger.info("Processing headers...")
                for section in document.sections:
                    header = section.header
                    for para in header.paragraphs:
                        if para.text.strip():
                            original_text = para.text.strip()
                            logger.info(f"Found header text to translate: {original_text}")
                            
                            self.enhanced_document_processor.process_paragraph(para, translate)
                            
                            translated_text = para.text.strip()
                            if translated_text != original_text:
                                translation_count += 1
                            else:
                                unchanged_count += 1
            
            # Process footers if configured
            if self.translate_footers:
                logger.info("Processing footers...")
                for section in document.sections:
                    footer = section.footer
                    for para in footer.paragraphs:
                        if para.text.strip():
                            original_text = para.text.strip()
                            logger.info(f"Found footer text to translate: {original_text}")
                            
                            self.enhanced_document_processor.process_paragraph(para, translate)
                            
                            translated_text = para.text.strip()
                            if translated_text != original_text:
                                translation_count += 1
                            else:
                                unchanged_count += 1
            
            logger.info("\nTranslation Statistics:")
            logger.info(f"- Texts translated: {translation_count}")
            logger.info(f"- Texts unchanged: {unchanged_count}")
                        
            # Save translated document
            logger.info(f"\nSaving translated document to: {output_path}")
            document.save(output_path)
            logger.info("Translation completed successfully")
            
        except Exception as e:
            logger.error(f"Translation failed: {str(e)}")
            raise

def translate_word(
    input_path: str,
    output_path: str,
    api_key: str,
    model: str = "gpt-4o",
    translate_headers: bool = True,
    translate_footers: bool = True,
    translate_comments: bool = False
) -> None:
    """
    Convenience function to translate a Word file.
    
    Args:
        input_path: Path to the input Word file
        output_path: Path where the translated file will be saved
        api_key: OpenAI API key
        model: OpenAI model to use (default: gpt-4o-mini)
        translate_headers: Whether to translate headers (default: True)
        translate_footers: Whether to translate footers (default: True)
        translate_comments: Whether to translate comments (default: False)
    """
    translator = WordTranslator(
        api_key,
        model,
        translate_headers=translate_headers,
        translate_footers=translate_footers,
        translate_comments=translate_comments
    )
    translator.translate_document(input_path, output_path)

if __name__ == "__main__":
    # Example usage - set your API key in environment variable OPENAI_API_KEY
    import os
    API_KEY = os.getenv('OPENAI_API_KEY')
    
    if not API_KEY:
        print("Please set OPENAI_API_KEY environment variable")
        exit(1)
    
    # Example usage with test file
    translate_word(
        input_path='input.docx',
        output_path='output_translated.docx',
        api_key=API_KEY,
        translate_headers=True,
        translate_footers=True,
        translate_comments=False
    )