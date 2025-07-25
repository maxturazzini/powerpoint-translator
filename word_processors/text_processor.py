from typing import List, Tuple, Optional
from docx.text.paragraph import Paragraph
from docx.text.run import Run
import logging

logger = logging.getLogger('word_translator')

class WordTextProcessor:
    """Text processing utilities for Word documents"""
    
    @staticmethod
    def get_paragraph_content(paragraph: Paragraph) -> str:
        """Extract text content from a paragraph"""
        try:
            return paragraph.text
        except Exception as e:
            logger.error(f"Error extracting paragraph content: {str(e)}")
            return ""
    
    @staticmethod
    def get_run_content(run: Run) -> str:
        """Extract text content from a run"""
        try:
            return run.text
        except Exception as e:
            logger.error(f"Error extracting run content: {str(e)}")
            return ""
    
    @staticmethod
    def extract_document_content(document) -> List[Tuple[str, str]]:
        """
        Extract all text content from document with context identifiers.
        Returns list of (text, context_id) tuples.
        """
        content_list = []
        
        try:
            # Extract from main document paragraphs
            for para_idx, paragraph in enumerate(document.paragraphs):
                text = WordTextProcessor.get_paragraph_content(paragraph)
                if text.strip():
                    context_id = f"main_para_{para_idx}"
                    content_list.append((text, context_id))
            
            # Extract from tables
            for table_idx, table in enumerate(document.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            text = WordTextProcessor.get_paragraph_content(paragraph)
                            if text.strip():
                                context_id = f"table_{table_idx}_r{row_idx}_c{cell_idx}_p{para_idx}"
                                content_list.append((text, context_id))
            
            # Extract from headers
            for section_idx, section in enumerate(document.sections):
                header = section.header
                for para_idx, paragraph in enumerate(header.paragraphs):
                    text = WordTextProcessor.get_paragraph_content(paragraph)
                    if text.strip():
                        context_id = f"header_s{section_idx}_p{para_idx}"
                        content_list.append((text, context_id))
            
            # Extract from footers
            for section_idx, section in enumerate(document.sections):
                footer = section.footer
                for para_idx, paragraph in enumerate(footer.paragraphs):
                    text = WordTextProcessor.get_paragraph_content(paragraph)
                    if text.strip():
                        context_id = f"footer_s{section_idx}_p{para_idx}"
                        content_list.append((text, context_id))
        
        except Exception as e:
            logger.error(f"Error extracting document content: {str(e)}")
        
        return content_list
    
    @staticmethod
    def count_runs_in_paragraph(paragraph: Paragraph) -> int:
        """Count the number of runs in a paragraph"""
        try:
            return len(paragraph.runs)
        except Exception as e:
            logger.error(f"Error counting runs: {str(e)}")
            return 0
    
    @staticmethod
    def has_complex_formatting(paragraph: Paragraph) -> bool:
        """Check if paragraph has complex formatting (multiple runs with different styles)"""
        try:
            if len(paragraph.runs) <= 1:
                return False
            
            # Check if runs have different formatting
            first_run = paragraph.runs[0]
            first_font = first_run.font
            
            for run in paragraph.runs[1:]:
                if (run.font.bold != first_font.bold or 
                    run.font.italic != first_font.italic or
                    run.font.underline != first_font.underline or
                    run.font.name != first_font.name):
                    return True
            
            return False
        except Exception as e:
            logger.error(f"Error checking complex formatting: {str(e)}")
            return False