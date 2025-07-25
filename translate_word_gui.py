import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import threading
import webbrowser
import platform
import subprocess
from translate_word import WordTranslator, WORD_INSTRUCTIONS
from docx import Document
from word_processors.text_processor import WordTextProcessor


def export_document_to_markdown(docx_path: str):
    """Export all content of a Word document to a Markdown file."""
    doc = Document(docx_path)
    md_lines = []
    summary = []
    
    # Process main document paragraphs
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = WordTextProcessor.get_paragraph_content(paragraph)
        if text.strip():
            # Determine heading level based on style
            style_name = paragraph.style.name if paragraph.style else ""
            if "Heading" in style_name:
                level = 1
                try:
                    level = int(style_name.split()[-1])
                except:
                    level = 1
                md_lines.append(f"{'#' * level} {text}")
            else:
                md_lines.append(text)
            md_lines.append("")
            summary.append(f"{os.path.basename(docx_path)} | Para {idx} | {text[:50]}...")
    
    # Process tables
    for table_idx, table in enumerate(doc.tables, 1):
        md_lines.append(f"## Table {table_idx}")
        
        # Add table header
        if table.rows:
            header_cells = []
            for cell in table.rows[0].cells:
                header_cells.append(cell.text.strip() or " ")
            md_lines.append("| " + " | ".join(header_cells) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
            
            # Add table rows
            for row in table.rows[1:]:
                row_cells = []
                for cell in row.cells:
                    row_cells.append(cell.text.strip() or " ")
                md_lines.append("| " + " | ".join(row_cells) + " |")
        
        md_lines.append("")
    
    md_path = os.path.splitext(docx_path)[0] + ".md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\\n".join(md_lines))
    
    return summary

# Load OpenAI API key from .env using python-dotenv
try:
    from dotenv import load_dotenv
    # Look for .env in current directory first, then parent directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    env_paths = [
        os.path.join(current_dir, '.env'),  # Current directory
        os.path.join(os.path.dirname(current_dir), '.env')  # Parent directory
    ]
    
    for env_path in env_paths:
        if os.path.exists(env_path):
            load_dotenv(env_path)
            break
    
    # Try different possible environment variable names
    ENV_API_KEY = os.getenv('openai_api_key') or os.getenv('OPENAI_API_KEY')
except ImportError:
    ENV_API_KEY = None

class WordTranslatorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Document Translator GUI")
        self.api_key = ENV_API_KEY or ""
        self.model = "gpt-4o"
        self.translated_path = None

        # File selection
        self.input_path_var = tk.StringVar()
        self.output_path_var = tk.StringVar()
        self.prompt_var = tk.StringVar(value=WORD_INSTRUCTIONS.strip())

        tk.Label(root, text="Select Word document to translate:").pack(anchor="w", padx=10, pady=(10,0))
        file_frame = tk.Frame(root)
        file_frame.pack(fill="x", padx=10)
        tk.Entry(file_frame, textvariable=self.input_path_var, width=60).pack(side="left", expand=True, fill="x")
        tk.Button(file_frame, text="Browse", command=self.browse_file).pack(side="left", padx=5)

        tk.Label(root, text="Destination file (auto-generated):").pack(anchor="w", padx=10, pady=(10,0))
        out_frame = tk.Frame(root)
        out_frame.pack(fill="x", padx=10)
        out_entry = tk.Entry(out_frame, textvariable=self.output_path_var, width=60, state="readonly")
        out_entry.pack(side="left", expand=True, fill="x")

        tk.Label(root, text="System Prompt (edit if needed):").pack(anchor="w", padx=10, pady=(10,0))
        self.prompt_text = scrolledtext.ScrolledText(root, height=12, wrap=tk.WORD)
        self.prompt_text.pack(fill="both", expand=True, padx=10)
        self.prompt_text.insert(tk.END, WORD_INSTRUCTIONS.strip())

        # Options frame
        options_frame = tk.Frame(root)
        options_frame.pack(fill="x", padx=10, pady=5)
        
        self.translate_headers_var = tk.BooleanVar(value=True)
        self.translate_footers_var = tk.BooleanVar(value=True)
        self.translate_comments_var = tk.BooleanVar(value=False)
        
        tk.Checkbutton(options_frame, text="Translate Headers", 
                      variable=self.translate_headers_var).pack(side="left")
        tk.Checkbutton(options_frame, text="Translate Footers", 
                      variable=self.translate_footers_var).pack(side="left", padx=10)
        tk.Checkbutton(options_frame, text="Translate Comments", 
                      variable=self.translate_comments_var).pack(side="left")

        # Translate button
        self.translate_btn = tk.Button(root, text="Translate", command=self.start_translation)
        self.translate_btn.pack(pady=10)

        # Open file button (hidden by default)
        self.open_btn = tk.Button(root, text="Open Translated File", command=self.open_translated_file)
        self.open_btn.pack(pady=(0,10))
        self.open_btn.pack_forget()

        # Export to Markdown
        self.export_btn = tk.Button(root, text="Export Document to Markdown", command=self.export_document)
        self.export_btn.pack(pady=(0,10))

        # Status
        self.status_var = tk.StringVar()
        
        # Set initial status based on API key availability
        if self.api_key and self.api_key.strip():
            self.status_var.set("✅ API key loaded - Ready to translate")
        else:
            self.status_var.set("⚠️ No API key found - Will prompt during translation")
            
        tk.Label(root, textvariable=self.status_var, fg="blue").pack(pady=(0,10))

    def browse_file(self):
        # Ask user to choose between file or folder
        choice = messagebox.askyesno("Selection Type", "Select 'Yes' for a single file or 'No' for a folder")
        
        if choice:  # Single file
            path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
            if path:
                self.input_path_var.set(path)
                base, ext = os.path.splitext(path)
                dest = f"{base}_Translated{ext}"
                self.output_path_var.set(dest)
                self.translated_path = dest
                self.open_btn.pack_forget()
        else:  # Folder
            path = filedialog.askdirectory()
            if path:
                self.input_path_var.set(path)
                dest = f"{path}_Translated"
                self.output_path_var.set(dest)
                self.translated_path = dest
                self.open_btn.pack_forget()

    def start_translation(self):
        input_path = self.input_path_var.get()
        output_path = self.output_path_var.get()
        prompt = self.prompt_text.get("1.0", tk.END).strip()
        
        if not input_path or not output_path:
            messagebox.showerror("Error", "Please select a Word document.")
            return
        if not self.api_key or not self.api_key.strip():
            self.api_key = self.ask_api_key()
            if not self.api_key or not self.api_key.strip():
                return
        
        self.status_var.set("Translating... Please wait.")
        self.translate_btn.config(state=tk.DISABLED)
        self.open_btn.pack_forget()
        
        # Get translation options
        translate_headers = self.translate_headers_var.get()
        translate_footers = self.translate_footers_var.get()
        translate_comments = self.translate_comments_var.get()
        
        threading.Thread(
            target=self.run_translation, 
            args=(input_path, output_path, prompt, translate_headers, translate_footers, translate_comments), 
            daemon=True
        ).start()

    def ask_api_key(self):
        # Simple dialog for API key
        return simple_input_dialog(self.root, "Enter OpenAI API Key:")

    def run_translation(self, input_path, output_path, prompt, translate_headers, translate_footers, translate_comments):
        try:
            # Check if input is a file or folder
            if os.path.isfile(input_path):
                # Single file translation
                translator = WordTranslator(
                    api_key=self.api_key,
                    model=self.model,
                    translate_headers=translate_headers,
                    translate_footers=translate_footers,
                    translate_comments=translate_comments
                )
                
                # Patch the instructions for this run
                global WORD_INSTRUCTIONS
                WORD_INSTRUCTIONS = prompt
                
                translator.translate_document(input_path, output_path)
                self.status_var.set(f"Done! Saved to: {output_path}")
                self.translated_path = output_path
                self.open_btn.pack(pady=(0,10))
                messagebox.showinfo("Success", f"Translation complete!\\nSaved to: {output_path}")
                
            else:
                # Folder translation
                if not os.path.exists(output_path):
                    os.makedirs(output_path)
                
                docx_files = [f for f in os.listdir(input_path) if f.lower().endswith('.docx')]
                if not docx_files:
                    messagebox.showinfo("No files", "No Word documents found in the selected folder.")
                    return
                
                translated_count = 0
                for filename in docx_files:
                    try:
                        input_file = os.path.join(input_path, filename)
                        base_name = os.path.splitext(filename)[0]
                        output_file = os.path.join(output_path, f"{base_name}_Translated.docx")
                        
                        self.status_var.set(f"Translating {filename}...")
                        
                        translator = WordTranslator(
                            api_key=self.api_key,
                            model=self.model,
                            translate_headers=translate_headers,
                            translate_footers=translate_footers,
                            translate_comments=translate_comments
                        )
                        
                        translator.translate_document(input_file, output_file)
                        translated_count += 1
                        
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to translate {filename}: {str(e)}")
                
                self.status_var.set(f"Done! Translated {translated_count} files to: {output_path}")
                self.translated_path = output_path
                self.open_btn.pack(pady=(0,10))
                messagebox.showinfo("Success", 
                                  f"Translation complete!\\nTranslated {translated_count} files\\nSaved to: {output_path}")
                
        except Exception as e:
            self.status_var.set(f"Error: {e}")
            messagebox.showerror("Error", str(e))
        finally:
            self.translate_btn.config(state=tk.NORMAL)

    def open_translated_file(self):
        if self.translated_path and os.path.exists(self.translated_path):
            system = platform.system()
            path = os.path.abspath(self.translated_path)
            if system == "Darwin":
                subprocess.run(["open", "-R", path])
            elif system == "Windows":
                os.startfile(os.path.normpath(path))
            else:
                subprocess.run(["xdg-open", path])
        else:
            messagebox.showerror("Error", "Translated file not found.")

    def export_document(self):
        """Export Word document to Markdown from a file or folder."""
        file_path = filedialog.askopenfilename(title="Select Word document", filetypes=[("Word files", "*.docx")])
        paths = []
        if file_path:
            paths = [file_path]
        else:
            folder = filedialog.askdirectory(title="Select folder containing DOCX files")
            if not folder:
                return
            for name in os.listdir(folder):
                if name.lower().endswith(".docx"):
                    paths.append(os.path.join(folder, name))

        if not paths:
            messagebox.showinfo("No files", "No Word documents selected.")
            return

        summary_lines = []
        for p in paths:
            self.status_var.set(f"Exporting {os.path.basename(p)}")
            self.root.update_idletasks()
            try:
                summary_lines.extend(export_document_to_markdown(p))
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export {p}: {e}")

        if summary_lines:
            if len(paths) > 1:
                # Multiple files - create summary
                folder = os.path.dirname(paths[0])
                md_path = os.path.join(folder, "AllTheDocuments.md")
                
                with open(md_path, "w", encoding="utf-8") as f:
                    for line in summary_lines:
                        f.write(line + "\\n")
        
        self.status_var.set("Markdown export completed")

def simple_input_dialog(parent, prompt):
    dialog = tk.Toplevel(parent)
    dialog.title("Input Required")
    tk.Label(dialog, text=prompt).pack(padx=10, pady=10)
    entry = tk.Entry(dialog, width=50, show="*")
    entry.pack(padx=10, pady=5)
    entry.focus_set()
    result = []
    def on_ok():
        result.append(entry.get())
        dialog.destroy()
    tk.Button(dialog, text="OK", command=on_ok).pack(pady=10)
    dialog.transient(parent)
    dialog.grab_set()
    parent.wait_window(dialog)
    return result[0] if result else None

if __name__ == "__main__":
    root = tk.Tk()
    app = WordTranslatorGUI(root)
    root.mainloop()